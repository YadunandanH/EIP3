# -*- coding: utf-8 -*-
"""
Created on Wed Jul 25 15:20:22 2018

@author: yadunandan.huded
"""

from bs4 import BeautifulSoup
from nltk.util import ngrams
import nltk 
from nltk import ne_chunk, pos_tag,word_tokenize
import pandas as pd
import itertools
import string
import re

from sklearn.metrics import confusion_matrix
from sklearn.ensemble import RandomForestClassifier

from collections import defaultdict
from sklearn import preprocessing

#Parsing html page to fetch addresses
soup = BeautifulSoup(open("C:\\Users\\yadunandan.huded\\Desktop\\Adress_classification\\storefinder_02.html", encoding="utf8"), "html.parser")

address_list = []
for stores in soup.find_all('div', {'id' : 'storeList'}):
    for stores_adress in stores.find_all(lambda tag: tag.name == 'p' and not tag.attrs):
        
        address_list.append(' '.join(stores_adress.text.replace('\n',' ').replace('\t',' ').split()))


#Steps to generate ngrams
#1. Tokenize the address
#2. Generate the Trigrams

tokenized_address = []
trigrams_list = []
pos_tagged_list = []

for i in range(0,len(address_list)):
    tokenized_address.append(list(nltk.word_tokenize(address_list[i])))
    trigrams_list.append(list(ngrams(tokenized_address[i],3)))
    for j in range(0,len(list(ngrams(tokenized_address[i],3)))):
        pos_tagged_list.append(nltk.pos_tag(list(ngrams(tokenized_address[i],3))[j]))

columns_list = ['ngram','ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','POS1','POS2','POS3','Contains_Digit','DigitLen','ISGeoNext','target']
df_features = pd.DataFrame(columns=columns_list)

#flatten list
flatten_list = list(itertools.chain.from_iterable(trigrams_list))

df_features['ngram'] = flatten_list

from nltk import Tree

def get_continuous_chunks(text, label):
    chunked = ne_chunk(pos_tag(word_tokenize(text)))
    print(chunked)
    prev = None
    continuous_chunk = []
    current_chunk = []

    for subtree in chunked:
        print("subtreee>>>>>>>>>>>>>>>>>>>>>>")
        print(subtree)
        print(type(subtree))
        print(subtree.label())
        if type(subtree) == Tree and subtree.label() in label:
            print("Entered")
            print(subtree.leaves())
            current_chunk.append(" ".join([token for token, pos in subtree.leaves()]))
            print(current_chunk)
        
    return current_chunk

#Loc_string = "WASHINGTON -- In the wake of a string of abuses by New York police officers in the 1990s, Loretta E. Lynch, the top federal prosecutor in Brooklyn, spoke forcefully about the pain of a broken trust that African-Americans felt and said the responsibility for repairing generations of miscommunication and mistrust fell to law enforcement."
Loc_string = ' '.join(df_features['ngram'][19])
NE = get_continuous_chunks(Loc_string, ['GPE','GSP'])


def feature_set(df):
    
    for i in range(0,df.shape[0]):
        term1,term2,term3,term4,term5 = '','','','',''
        digilen = 0
        Geo_len = 0
        for j in range(0,3):
            if df['ngram'][i][j][0].isupper():
                term1 +='P'
            else:
                term1 +='N'
            if df['ngram'][i][j].islower():
                term2 +='P'
            else:
                term2 +='N'
            if df['ngram'][i][j][0].isnumeric():
                term3 += 'P'
            else:
                term3 += 'N' 
            if set(list(df['ngram'][i][j])).intersection(set(string.punctuation)):
                term4 += 'P'
            else:
                term4 += 'N'
            if bool(re.search(r'\d',df['ngram'][i][j])):
                term5 += 'P'
            else:
                term5 += 'N'
                         
                
        df['ISCAPITAL'][i] =  term1
        df['ISLOWER'][i] = term2
        df['ISDIGIT'][i] = term3
        df['Puncuation'][i] = term4        
        df['Contains_Digit'][i] = term5 

        if df['ngram'][i][0].isnumeric():
                digilen += len(df['ngram'][i][0])
                Loc_string = ' '.join(df['ngram'][i])
                NE = get_continuous_chunks(Loc_string, ['GPE','GSP'])
                #print(NE)
                Geo_len = len(NE)
                print(Geo_len)

        df['DigitLen'][i] = digilen
        if Geo_len:
            print('entered**********')
            df['ISGeoNext'][i] = 1
        else:
             df['ISGeoNext'][i] = 0
        
        pos_tags = nltk.pos_tag(nltk.word_tokenize(' '.join(df['ngram'][i])))
        #print(pos_tags[0][1])
        
        df['POS1'][i] = pos_tags[0][1]
        df['POS2'][i] = pos_tags[1][1]
        df['POS3'][i] = pos_tags[2][1]
    
    return df
 

df_features = feature_set(df_features)

df_features.to_excel(r'Feature_set2.xlsx')

#data = pd.read_excel(r'Address_Data_1st_cut.xlsx')

data.columns


d = defaultdict(preprocessing.LabelEncoder)
# Encoding the variable
fit = data[['ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit']].apply(lambda x: d[x.name].fit_transform(x))
'''
# Inverse the encoded
fit.apply(lambda x: d[x.name].inverse_transform(x))

# Using the dictionary to label future data
df.apply(lambda x: d[x.name].transform(x))
'''
Model_data = data.copy()

Model_data[['ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit']] = fit[['ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit']].copy()

y = Model_data['target']
X = Model_data.drop(['ngram','target'],axis = 1)



clf = RandomForestClassifier(n_estimators=40,max_depth=8,max_features=.5, random_state=56,min_samples_leaf=5,class_weight='balanced')

clf.fit(X, y)

print(clf.feature_importances_)

pred = clf.predict(X)

confusion_matrix(y,pred)

test_data = pd.read_excel(r'C:\Users\yadunandan.huded\Desktop\Test_address.xlsx')

test_fit = test_data[['ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit']].apply(lambda x: d[x.name].fit_transform(x))
test_data[['ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit']] = test_fit[['ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit']]
 
cols = ['ISCountry','ISCAPITAL','ISLOWER','ISDIGIT','Puncuation','PrevPOS','POS1','POS2','POS3','Contains_Digit','DigitLen']

pred_test = clf.predict(test_data[cols])

confusion_matrix(test_data['target'],pred_test)  


#Import document
import docx

doc = docx.Document(r'C:\Users\yadunandan.huded\Desktop\Backup03282019\JULY_Cutoffs.docx')
len(doc.paragraphs)
full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)
